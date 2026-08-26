import os
from docx import Document
from docxcompose.composer import Composer

def test_compose():
    # 1. Create a blank master doc and add something
    master = Document()
    master.add_heading('Hoja de Ruta')
    
    composer = Composer(master)
    
    # 2. Append a filled template
    template_path = r'C:\CRM PYP\plantilla_caja_huancayo_v2.docx'
    doc_temp = Document(template_path)
    
    # Simulate filling template
    for p in doc_temp.paragraphs:
        if '[NOMBRE_CLIENTE]' in p.text:
            p.text = p.text.replace('[NOMBRE_CLIENTE]', 'JUAN PEREZ')
            
    # Save to memory or temp file? docxcompose requires a Document object. 
    # But does it require a saved document? No, you just pass the Document object.
    try:
        composer.append(doc_temp)
        composer.save(r'C:\CRM PYP\test_compose_output.docx')
        print("Success!")
    except Exception as e:
        print("Error:", e)

if __name__ == '__main__':
    test_compose()
