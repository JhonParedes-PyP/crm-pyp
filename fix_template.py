import docx
import os

doc_path = r'C:\CRM PYP\MODELO DE CARTAS CAJA HUANCAYO 26.docx'
out_path = r'C:\CRM PYP\plantilla_caja_huancayo_v2.docx'

doc = docx.Document(doc_path)

mapping = {
    'K & P MAQUINARIAS S.A.C': '[NOMBRE_CLIENTE]',
    '107099101001382366': '[NUM_CUENTA]',
    'MARTINEZ PALOMARES, CESAR PAUL': '[NOMBRE_AVAL]',
    'AG. CHORRILLOS': '[AGENCIA]',
    'AA.HH. BUENOS AIRES  MZ. B2 LT. 10 SAN JUAN DE MIRAFLORES': '[DIRECCION_CLIENTE]',
    'AA.HH. BUENOS AIRES  MZ. B2 LT. 10 POR MIGUEL IGLESIAS AA.HH 13 DE OCTUBRE SAN JUAN DE MIRAFLORES': '[DIRECCION_CLIENTE]',
    '36325.52': '[MONTO_DEUDA]',
    '1997-2026-CÓDIGO AG. 099 ASESOR LEGAL EXTERNO: LPG/CMACHYO': '[NRO_CARTA]',
    'FECHA ULTIMO PAGO: / /2026': 'FECHA ULTIMO PAGO: [FECHA_ULT_PAGO]'
}

def replace_in_paragraph(paragraph, mapping):
    for key, val in mapping.items():
        if key in paragraph.text:
            replaced_in_run = False
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, str(val))
                    replaced_in_run = True
            
            if not replaced_in_run:
                # Spans multiple runs. Keep first run's formatting for the whole paragraph to be safe
                full_text = paragraph.text.replace(key, str(val))
                if paragraph.runs:
                    paragraph.runs[0].text = full_text
                    for r in paragraph.runs[1:]:
                        r.text = ""

for p in doc.paragraphs:
    replace_in_paragraph(p, mapping)

for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, mapping)

# Add page break before CARGO
p_idx = -1
for i, p in enumerate(doc.paragraphs): 
    if 'CARGO' in p.text and len(p.text) < 15: 
        p_idx = i; break;
if p_idx != -1: 
    doc.paragraphs[p_idx].insert_paragraph_before('').add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

# Page setup: margins to 0.5 cm (approx 14.17 points)
from docx.shared import Cm
for section in doc.sections:
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)

doc.save(out_path)
print("Template saved to", out_path)
