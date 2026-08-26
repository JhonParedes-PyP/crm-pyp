from docx import Document

def test_clear():
    doc = Document(r'C:\CRM PYP\plantilla_caja_huancayo_v2.docx')
    
    # Remove all elements
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    for t in doc.tables:
        t._element.getparent().remove(t._element)
        
    # Word requires at least one paragraph
    doc.add_paragraph()
    
    doc.save(r'C:\CRM PYP\test_blank.docx')
    print("Success")

if __name__ == '__main__':
    test_clear()
